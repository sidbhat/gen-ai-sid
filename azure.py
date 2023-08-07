import streamlit_analytics
from open_ai_service import OpenAIService
import streamlit as st
from io import BytesIO
import docx
import datetime
import pinecone
from langchain.chains.question_answering import load_qa_chain
from langchain.text_splitter import CharacterTextSplitter, RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma, Pinecone
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.llms import OpenAIChat
import os
import spacy
import nlp
import pandas as pd
# from googlesearch import search
import requests
from bs4 import BeautifulSoup
# from PyPDF2 import PdfReader
import time

pd.set_option("max_colwidth", 300)
nlp = spacy.load("en_core_web_sm")
ruler = nlp.add_pipe("entity_ruler")
patterns = [{"label": "ORG", "pattern": "Eightfold"}, {"label": "ORG", "pattern": "Successfactors"},
            {"label": "ORG", "pattern": "SAP"},
            {"label": "GPE", "pattern": [{"LOWER": "san"}, {"LOWER": "francisco"}]}]
ruler.add_patterns(patterns)

sap_options1 = ["", "Tell me the features of mobile time recording for iOS and Android",
                "Provide the key highlights from Time Management 1H 2023 Release Notes",
                "What are the new features of the Weekly Time Sheet",
                "What are the features available for grace rule rounding and rest-rules",
                "How do we position our solution offerings in the Time Management space?",
                "What is our Time Management Solution Extensions Strategy",
                "What are the main features of the 1H 2023 release of the SAP SuccessFactors HXM Suite",
                "How does the SAP SuccessFactors Opportunity Marketplace help employees",
                "How can I get ready for Talent Intelligence Hub?",
                "How does Talent Intelligence Hub interact with other SAP SuccessFactors modules?",
                "What is the SAP SuccessFactors Innovation Strategy?",
                "What are some examples of SAP Solution Extensions",
                "What is Eightfold and how does it integrate with SAP?",
                "What is Beamery and how does it integrate with SAP?",
                "What is the learning plan to be become a Successfactors LMS consultant? What skills do I need?",
                "What are the different types of partnerships that SAP offers? What is a solution extension partner?",
                "How do I change the theme in the Successfactors application? Provide steps for an adminstrator",
                "You are an HR industry expert. Provide the top HR challenges facing the retail industry",
                "Provide 10 limbic openings for a presentation about talent management, learning and payroll.",
                "Provide in a table format an Employee table with the following columns and 10 randomized entries \n [EmployeeId, Employee Name, Job Classification, Cost Center, Region, Job Location, Department, Average Tenure, Total YOE, Pay Grade, Total CTC, Compa- ratio. Impact-of-leaving, Cost-to-train/year, Performance Rating, Future Leader]"]
sap_options2 = ["", "You are an HR industry expert. Provide the top HR challenges facing the retail industry",
                "Write a 300-word SEO-friendly blog with right keywords, title, meta-description, h1,h2 details and  5 examples on how the gig economy will change the retail industry.",
                "Write an outline for an ebook with 10 chapters on how the future of hybrid work can be improved by Web3 and Metaverse. Provide detailed titles and sites as references.",
                "Provide a short blog post on job functions that may be displaced by AI in the financial sector with an intro, description and conclusion section"]
sap_options3 = ["",
                "Rewrite this demo script for a Retail customer  \n Introduction: \n HXM, Human Experience Management, is an evolution of HCM. HXM is about creating individualized, dynamic experiences for all users - whether candidates, new hires, employees, managers or HR leaders.Because our people determine whether an organization succeeds or fails, the experiences we deliver and how we address the moments that matter for every individual impact everything from top line business growth to profitability to how quickly an organization can adapt to change.With HXM we enable customers to move beyond simply facilitating transactions and supporting HR-driven processes, to focus on creating a human-centered approach.\nThis means putting your people and their experiences at the center of everything, listening to their needs and providing them with the tools and information to be productive and make a real impact on the business. Our HXM Suite helps you redefine experiences -- making everyone (candidates, new hires, employees, managers, and HR leaders) feel connected, engaged, empowered, and supported at every step in their employee journey. \nAbout This Demo Script:\nThis script follows one persona, Alex Smith, through the journey from candidate to new hire to employee and, finally, to manager.While the story assumes that you will show Alex as the same persona throughout the demo, you may prefer to show different personas. If so, you will need to adjust the accompanying slide deck and talk track accordingly.\nVignette 1 \n- Intro to Candidate ExperienceAs a candidate, you ask yourself: Can you imagine yourself working for this organization?\n You are driven by meaningful work, supportive management, fantastic work environment, and growth opportunities amongst other things. Will this Company offer that to you? \n You want to understand this company’s culture: does their Corporate Culture align with your priorities? You want to get a sense of their culture including work/life balance, purpose, diversity, etc. through the recruitment experience.\nDemo Flow Demo Story/Message\n. What we want to highlight through this story and solution demonstration is how human experience management (HXM) helps organizations go beyond simply facilitating transactions to truly reinvent experiences in ways that accelerate business growth and drive engagement. \n We do this by shifting the focus from supporting top-down, HR-driven processes to delivering individualized experiences designed completely around the needs of the people in your organization.As we think about how you are going to meet the expectations of your workforce, we need to think about what the ideal experience is for them. \n In fact, we need to start with the experience before they even become an employee. What exactly are candidates looking for?When you think about critical moments for a candidate – searching for a job, the application process, multiple interviews, etc., each of those typically result in many questions and uncertainties. \n  But many organizations offer a less than engaging candidate experience.When it comes to attracting the right talent and offering an engaging candidate experience, many companies struggle with challenges such as a requisition-centric process, high cost of candidate acquisition, long time to hire, and competing with other organizations trying to attract the same high-quality candidates.Organizations face a continued need to attract workers with technological and other specialized skills, yet many have an inability to effectively convey their employer brand.\n Poor sourcing methodologies often lead to ineffective talent pools. Many have antiquated recruiting systems that tend to be limited to applicant tracking and lack complete applicant marketing and mobile capabilities.\n With the current talent shortages, we’re all reading about, this leaves many organizations unable to compete in what some are now calling a new “War for Talent”.",
                "Provide a audio script track for the SuccessFactors recruiting solution for a fashion retail business. Include a headline, intro, description and outro",
                "Provide 10 limbic openings for a presentation about talent management, learning and payroll.",
                "Provide a talk track introducing Successfactors HR to retail customers with HR leaders in the audience",
                "Write an email talking about the perennial HR challenges with a thematic tie-in to the World Cup finals.",
                "You are a presales expert. Write an email to [person] with some facts about how SuccessFactors has over 3000 EC Customers with a thematic tie-in to Christmas.",
                "Write an email talking about the perennial HR challenges around hiring, motiving and guiding employees in English and German",
                "Generate a sample Non Disclosure Agreement document for retail company 'Best Run' for its external contractors",
                "Help me generate a AI image prompt similar to - Award-winning mix-use complex contemporary designed by the best architects in England, stunning London riverside, high resolution, ultra-detailed, 8k, architectural photography Archdaily, Hyper-realistic, intricate detail, photorealistic",
                "Write a twitter thread on the benefits of blockchain based verified employee credentials",
                "Craft an SEO friendly Linkedin post about diversity and inclusion in the workplace. Provide the target keywords."]
sap_options4 = ["",
                "Please generate a job description for a [Digital Marketing Specialist].The ideal candidate should have skills in [SEO/SEM, marketing database, email, social media, and display advertising campaigns]. \n They must be able to have experience in [leading integrated digital marketing campaigns from concept to execution] and have additional experience in [ knowledge of website analytics tools]. Please include the job responsibilities and required qualifications.Ensure the job description does not have bias and is inclusive.",
                "Write me python and abap code to make a REST API call and authenticate via Microsoft Active Directory",
                "Provide in a table format an Employee table with the following columns and 10 randomized entries \n [EmployeeId, Employee Name, Job Classification, Cost Center, Region, Job Location, Department, Average Tenure, Total YOE, Pay Grade, Total CTC, Compa- ratio. Impact-of-leaving, Cost-to-train/year, Performance Rating, Future Leader]",
                "Write me a VBA macro to create a presentation for a startup.",
                "Generate a sample google sheet with sample movies dataset that can be used for exploratory analysis."]

next_action = ["","Summarize this response highlighting the key takeaways.", "Expand on this response to provide more details.", "Explain this response like I am five."]
count_str = ""
result_str = ""
response = ""
st.set_page_config(page_title="Chat GPT| Open AI| SAP Generative AI| Sid Bhattacharya", page_icon=':rocket:',
                   layout='wide')
c = st.container()
header = c.container()
# loading_placeholder = c.container()

input_container = c.container()
output_container = c.container()
header.header("🚀 Ask Chatty McChatface v1")
header.caption(
    "Demo app that showcases how to do prompt engineering with Open AI and how to build an enterprise knowledge base using custom embeddings.")
bar = header.empty()

with header:
    with st.expander("💬 Features and Help"):
        st.write('''
    Features :
    1. **Prompt Engineering** - Primed prompts to provide factual information and cite references. 
    2. **Prompt List** - Out of the box 20+ prompts based on job functions and roles.
    2. **Enterprise Search** - Link to an enterprise knowledge base. Shows how domain and up to date information can be provided (Example - What is new in the Successfactors 2023 1H Release) .
    3. **Hide PII and Confidential Information** - Redact feature to hide PII data and other confidential information.
    4. **Save Results** - Export responses to a Word Document.
    5. **History** - View history of requests made.
    6. **New Prompts** - To showcase enterprise search available under 'Customers and Partners' drop down.  New prompt to showcase demo-scripting under 'Content & Marketing'.
    ***
    Prompts : 

    🕳 **Customers and Partners**
    1. What was <<enter company>> revenue in 2020.
    2. Compare <<enter company>> and <<enter company>> against Successfactors in the North America region.
    3. List HR software companies in Middle East and North Africa region listing their strengths in the region.
    4. What is the learning plan to be become a Successfactors LMS consultant? What skills do I need?
    5. What are the different types of partnerships that SAP offers? What is a solution extension partner?
    6. How do I change the theme in the Successfactors application? Provide steps for an adminstrator.
    7. You are HustleGPT, an entrepreneurial AI. You have $200, and your goal is to turn that into as much money as possible in the shortest time possible without doing anything illegal. Provide detailed steps.
    ***                 
    💡 **Industry & Emerging Trends**
    1. You are an HR industry expert. Provide the top HR challenges facing the retail industry.
    2. Write a 300-word SEO-friendly blog with right keywords, title, meta-description, h1,h2 details and  5 examples on how the gig economy will change the retail industry.
    3. Provide a short blog post on job functions that may be displaced by AI in the financial sector with an intro, description and conclusion section.
    4.Write an outline for an ebook with 10 chapters on how the future of hybrid work can be improved by Web3 and Metaverse. Provide detailed titles and sites as references.
    ***
    🗯 **Content Generation (Sales & Marketing)**
    1. Provide a audio script track for the SuccessFactors recruiting solution for a fashion retail business. Include a headline, intro, description and outro.
    2. Provide 10 limbic openings for a presentation about talent management, learning and payroll.
    3. Provide a talk track introducing Successfactors HR to retail customers with HR leaders in the audience.
    4. Write an email talking about the perennial HR challenges with a thematic tie-in to the World Cup finals.
    5. You are a presales expert. Write an email to [person] with some facts about how SuccessFactors has over 3000 EC Customers with a thematic tie-in to Christmas.
    6. Write an email talking about the perennial HR challenges around hiring, motiving and guiding employees in English and German.
    7. Help me generate a Midjourney prompt similar to - Award-winning mix-use complex contemporary designed by the best architects in England, stunning London riverside, high resolution, ultra-detailed, 8k, architectural photography Archdaily, Hyper-realistic, intricate detail, photorealistic
    8. Write a twitter thread on the benefits of blockchain based verified employee credentials
    ***
    💛 **Data & Code Generation**
    1. Please generate a job description for a [Digital Marketing Specialist].The ideal candidate should have skills in [SEO/SEM, marketing database, email, social media, and display advertising campaigns]. They must be able to have experience in [leading integrated digital marketing campaigns from concept to execution] and have additional experience in [ knowledge of website analytics tools]. Please include the job responsibilities and required qualifications.Ensure the job description does not have bias and is inclusive.
    2. Provide in a table format an Employee table with the following columns and 10 randomized entries EmployeeId, Employee Name, Job Classification, Cost Center, Region, Job Location, Department, Average Tenure, Total YOE, Pay Grade, Total CTC, Compa- ratio. Impact-of-leaving, Cost-to-train/year, Performance Rating, Future Leader.
    3. Write me python and abap code to make a REST API call and authenticate via Microsoft Active Directory.
    4. Write me a VBA macro to create a presentation for a startup.
    5. Generate a sample google sheet with sample movies dataset that can be used for exploratory analysis.
        ''')

sidebar_placeholder = st.sidebar.container()

if 'response' not in st.session_state:
    st.session_state.response = ''

if 'key' not in st.session_state:
    st.session_state['key'] = ''

if "conversations" not in st.session_state:
    st.session_state["conversations"] = ""

if "clear" not in st.session_state:
    st.session_state["clear"] = False

selected_value1 = ''
selected_value2 = ''
selected_value3 = ''
selected_value4 = ''
submitted = False

conversation = [{"role": "system",
                 "content": "You are a helpful assistant that provides detailed answers based on facts. Always cite references for your responses towards the end of the response. "}]



user_input = ''
search_results = ''
download_content = ''
index_name = 'demo-index'

@st.cache_resource
def init_model():
    return OpenAIEmbeddings(openai_api_key=os.environ['OPENAI_API_KEY'])

embeddings = init_model()

@st.cache_resource
def init_llm():
  return OpenAIChat(temperature=0, openai_api_key=os.environ['OPENAI_API_KEY'], model_name='gpt-3.5-turbo',
                 model_kwargs={'max_tokens': 1000})
llm = init_llm()

# def load_knowledge_base():
#     # Load PDFS
#     print("in knowledge base")
#     print(uploaded_file)
#     if uploaded_file:
#         pdf_reader = PdfReader(uploaded_file)
#         text = ""
#         for page in pdf_reader.pages:
#             text += page.extract_text()
#
#         # Create chunks
#         text_splitter = CharacterTextSplitter(
#             separator="\n",
#             chunk_size=1000,
#             chunk_overlap=200,
#             length_function=len
#         )
#         texts = text_splitter.split_text(text)
#
#         # loader = PyPDFLoader(uploaded_file.name)
#         # documents = loader.load()
#         # print(documents[:5])
#         # text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
#         # texts = text_splitter.split_documents(documents)
#         for t in texts:
#             print(t.capitalize())
#         Pinecone.from_texts([t.capitalize() for t in texts], embeddings, index_name=index_name)
#         sidebar_placeholder.write("Knowledge base updated")
#         sidebar_placeholder.caption(pinecone.Index(index_name).describe_index_stats())
#         # chain = load_qa_chain(llm, chain_type="stuff")
#         #
#         # query = "Summarize the document " + str_
#         #
#         # docsearch = Pinecone.from_existing_index(index_name, embeddings)
#         # docs = docsearch.similarity_search(query, include_metadata=True)
#         # print(len(docs))
#         # char_text_splitter = RecursiveCharacterTextSplitter(chunk_size=100, chunk_overlap=10)
#         # d = char_text_splitter.split_documents(docs)
#         # print(len(d))
#         # response = chain.run(input_documents=d[:1], question=query)
#         # print(response)
#         # print("done")
#
#
# # uploaded_file = sidebar_placeholder.file_uploader("🧠 Add to Knowledge Base", type=["PDF"])
# # if index_name in pinecone.list_indexes():
# #     sidebar_placeholder.caption(pinecone.Index(index_name).describe_index_stats())
# # sidebar_placeholder.button("Upload", on_click=load_knowledge_base(), disabled=True)

sidebar_placeholder.header('🍁 History')


@st.cache_resource
def init_pinecone():
    # initialize connection (get API key at app.pinecone.io)
    pinecone.init(
        api_key="5e6a8cb6-f036-4a23-9b34-c95aec8e317f",
        environment="us-west1-gcp-free"  # find next to API key
    )
    # connect to index
    if index_name not in pinecone.list_indexes():
        # if does not exist, create index
        pinecone.create_index(
            index_name,
            dimension=1536,  # dimensionality of text-embedding-ada-002
            metric='cosine',
        )
    index = pinecone.Index(index_name)
    print(index.describe_index_stats())


# def googlesearch(query, num_results=5):
#     search_results = list(search(query, num_results=num_results))
#     return search_results


def extract_search_result_info(link):
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3"
    }
    response = requests.get(link, headers=headers)
    soup = BeautifulSoup(response.text, "html.parser")

    title = soup.title.string.strip() if soup.title else "No title found"
    meta_description = soup.find("meta", attrs={"name": "description"})
    description = meta_description.get("content").strip() if meta_description else "No description found"
    excerpt = '\n'.join(description.split('\n')[:3]) if description else "No excerpt found"

    # Extract the publication date if available
    date = soup.find("meta", attrs={"property": "article:published_time"})
    date = date.get("content").strip() if date else "No date found"

    return title, excerpt, date


def replace_ner(mytxt):
    clean_text = mytxt
    doc = nlp(mytxt)
    for ent in reversed(doc.ents):
        clean_text = clean_text[:ent.start_char] + ent.label_ + clean_text[ent.end_char:]
    print("Redacted text" + clean_text)
    return clean_text


def enterprise_search(prompt_input: str):
    # connect to index
    docsearch = Pinecone.from_existing_index(index_name, embeddings)

    query = "You are a helpful assistant that provides detailed answers based on facts based on the given context. Provide your response in the format {question} {answer} followed by {references}. Always cite references for your responses towards the end of the response." + st.session_state.prompt
    chain = load_qa_chain(llm, chain_type="stuff")
    docs = docsearch.similarity_search(prompt_input)
    # print(docs)
    st.session_state.docs = docs
    st.session_state.enterprise_search = chain.run(input_documents=docs, question=query)
    # print(st.session_state.enterprise_search)
    # total_tokens = response.get("total_tokens")
    # # pricing logic: https://openai.com/pricing#language-models
    #     if st.session_state.model == "gpt-3.5-turbo":
    #         cost = total_tokens * 0.002 / 1000


# def save_to_pptx(ai_content: str):
#     prs = Presentation()
#     title_slide_layout = prs.slide_layouts[0]
#     slide = prs.slides.add_slide(title_slide_layout)
#     title = slide.shapes.title
#     subtitle = slide.placeholders[1]
#     # print(ai_content)
#     title.text = "Generated Content"
#     subtitle.text = "By Chatty McChatface"
#
#     # adding text
#     if st.session_state.response:
#         # for i in range(len(st.session_state.response)):
#         slide = prs.slides.add_slide(prs.slide_layouts[5])
#         title = slide.shapes.title
#         title.text = "Response from AI"
#         # For adjusting the  Margins in inches
#         left = top = width = height = Inches(1)
#
#         # creating textBox
#         txBox = slide.shapes.add_textbox(left - 0.5, top - 0.5,
#                                          width * 9, height * 7)
#         # creating textFrames
#         tf = txBox.text_frame
#         tf.word_wrap = True
#         tf.text = st.session_state.prompt
#         p = tf.add_paragraph()
#         p.font.size = Pt(14)
#         p.text = st.session_state.response
#         p.alignment = PP_PARAGRAPH_ALIGNMENT.LEFT
#
#     # save the output into binary form
#     binary_output = BytesIO()
#     prs.save(binary_output)
#
#     return binary_output.getvalue()
#
def reset_context():
    if st.session_state.get("clear"):
        bar.warning(":information_source: Conversation cleared....")
        time.sleep(2)
        bar.empty()
        del st.session_state["conversations"]
        st.session_state["prompt_list"] = sap_options1[0]
        st.session_state["prompt"] = ''
        st.session_state["bool_search"] = False
        st.session_state["redact"] = False
        st.session_state['key']=''
        selected_value1=""
        conversation=""


@st.cache_resource
def build_footer():
    c1, c2, c3 = c.columns(3)
    with c:
        with c1:
            st.info('**Contact Me: [@sidbhat5](https://twitter.com/sidbhat5)**', icon="💡")
        with c2:
            st.info('**Prompts Guide: [Prompts](https://sidbhat.blog/10-chat-gpt-prompts-to-help-you-succeed/)**',
                    icon="💻")
        with c3:
            st.info('**Google Collab: [Code](https://colab.research.google.com/drive/)**', icon="🧠")


def redact_string(str_: str):
    print(st.session_state.redact)
    if st.session_state.redact:
        return replace_ner(str_)
    else:
        return str_


def download_docx(str_: str):
    doc = docx.Document()

    # Add Title Page followed by section summary
    doc.add_heading("Generated Output", 0)
    doc.add_paragraph(f'Authored By: Chatty McChatface')

    doc.add_heading("Request", 1)
    doc.add_paragraph(st.session_state.prompt.capitalize())

    doc.add_heading("Response", 1)
    doc.add_paragraph(str_)
    binary_output = BytesIO()
    doc.save(binary_output)
    return binary_output.getvalue()



def send_click(prompt_input :str):
  with c:
      if prompt_input == '':
          bar.warning(":information_source: Enter a request...")
          time.sleep(2)
          bar.empty()
      else:
         # with loading_placeholder:
           with st.spinner("Fetching response..."):
              with output_container:
                    if st.session_state.conversations:
                        st.session_state.conversations.append({"role": "user", "content": prompt_input})
                    else:
                        ai_role = "You are a helpful SAP assistant that provides detailed answers based on facts. Provide your response in the format {question} {answer} followed by {references}. Always cite references for your responses towards the end of the response."  # NOQA: E501
                        st.session_state.conversations = [
                            {"role": "system", "content": ai_role},
                            {"role": "user", "content": prompt_input},
                        ]
                    conversation.append({"role": "user", "content":prompt_input})
                    st.session_state['key'] = st.session_state['key'] + '~~~' + prompt_input.capitalize()
                    if st.session_state['key']:
                       keys_list = st.session_state['key'].split('~~~')
                       keys_list_reversed = keys_list[::-1]
                       with sidebar_placeholder:
                            cnt=1
                            for keys in keys_list_reversed:
                                if (keys != ''):
                                   sidebar_placeholder.info(":thread: "+ str(cnt) +'/'+ str(len(keys_list)-1)+' '+keys)
                                   cnt = cnt+1

                    if st.session_state.bool_search:
                        c.caption(":eight_spoked_asterisk: Response from Enterprise Knowledge Base")
                        enterprise_search(prompt_input)
                        c.success(redact_string(st.session_state.enterprise_search))
                        st.session_state.conversations.append({"role": "assistant", "content": st.session_state.enterprise_search})
                        download_content = st.session_state.enterprise_search
                    else:
                        st.session_state.response = OpenAIService.open_ai_query(query='', model='gpt-4',
                                                                                gpt_conversation_history=st.session_state.conversations)
                        st.session_state.conversations.append({"role": "assistant", "content": st.session_state.response})
                        download_content = st.session_state.response
                        #print(conversation)
                        c.caption(":robot_face: Response from Open AI")
                        c.info(redact_string(st.session_state.response))

                    if download_content != '':
                        c.download_button(
                                    label="⬇️ Download",
                                    data=download_docx(redact_string(download_content)),
                                    file_name="Chatty McChatface Response-" + datetime.datetime.now().strftime(
                                        "%m-%d-%Y-%H:%M:%S") + ".docx",
                                    mime="docx"
                                )

            # build_footer()

init_pinecone()

def main():
        with streamlit_analytics.track():
            with input_container:
                st.session_state.clear = c.button(":negative_squared_cross_mark: Clear Conversation")
                if st.session_state.clear:
                    reset_context()
              # with st.form(key = 'my_form', clear_on_submit=True):
                selected_value1 = st.selectbox("💛 Prompts", sap_options1, key='prompt_list', index=0)
                if ( (selected_value1 != '' and "prompt" in st.session_state and st.session_state.prompt == '') or ("prompt" in st.session_state and st.session_state.prompt in sap_options1)):
                    st.session_state.prompt = selected_value1
                prompt_input = st.text_area("🦋 Ask something: ", key='prompt')
                if st.session_state.get("bool_search"):
                    bool_search = st.checkbox('Include Enterprise Knowledge Base', key='bool_search',
                                         value=st.session_state.bool_search)
                else:
                    st.session_state.bool_search = st.session_state.get("bool_search", False)
                    bool_search = st.checkbox('Include Enterprise Knowledge Base', key='bool_search',
                                         value=st.session_state.bool_search)
                if st.session_state.get("redact"):
                    st.checkbox('Redact Confidential Information', key="redact", value=st.session_state.redact)
                else:
                    st.session_state.redact = st.session_state.get("redact", False)
                    st.checkbox('Redact Confidential Information', key="redact", value=st.session_state.redact)

              #   col1, col2 = c.columns(2)
                # with col1:
                submitted = st.button("✅ Send Message")

                if submitted:
                      send_click(prompt_input)



if __name__ == "__main__":
        main()
